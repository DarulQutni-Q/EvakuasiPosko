import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def buat_makalah():
    doc = docx.Document()
    
    # Konfigurasi Font Default (Times New Roman, 12)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Buat Style untuk Heading (agar lebih mudah dan konsisten)
    def set_heading_style(level, size, bold=True):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Times New Roman'
        h_style.font.size = Pt(size)
        h_style.font.bold = bold
        h_style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
        
    set_heading_style(1, 14, True) # Bab
    set_heading_style(2, 12, True) # Sub-bab
    set_heading_style(3, 12, False) # Sub-sub-bab

    # Helper function untuk Numbering List agar indentasinya rapi
    def add_numbered_list(doc_obj, text):
        p = doc_obj.add_paragraph(text, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
        return p

    # --- HALAMAN JUDUL ---
    doc.add_paragraph('\n\n\n')
    judul = doc.add_paragraph()
    judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_judul = judul.add_run('LAPORAN PROYEK PEMROGRAMAN BERORIENTASI OBJEK\nSistem Manajemen Evakuasi Bencana dan Pengungsian\n\n')
    run_judul.font.size = Pt(14)
    run_judul.bold = True
    
    identitas = doc.add_paragraph()
    identitas.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_identitas = identitas.add_run('Disusun Oleh:\n\nMusyawirdi Mahtuah\nNIM: 220180153\n\n')
    run_identitas.bold = True
    
    doc.add_page_break()

    # --- BAB I ---
    doc.add_heading('BAB I: PENDAHULUAN', level=1)
    
    doc.add_heading('1.1 Analisis Masalah', level=2)
    doc.add_paragraph(
        "Dalam situasi tanggap darurat bencana, proses pendataan dan distribusi logistik kerap menjadi titik lemah. "
        "Pendataan manual sering kali menyebabkan duplikasi data warga, sementara pemantauan kapasitas posko yang tidak real-time memicu penumpukan pengungsi di satu titik. "
        "Selain itu, keterlambatan identifikasi kelompok rentan (seperti lansia dan anak-anak) menyebabkan distribusi bantuan medis dan logistik menjadi tidak tepat sasaran. "
        "Oleh karena itu, diperlukan sebuah sistem terpusat yang mampu mengotomatisasi penentuan prioritas evakuasi serta memberikan peringatan dini terkait daya tampung dan logistik posko secara otomatis."
    )
    
    doc.add_heading('1.2 Kebutuhan Sistem', level=2)
    doc.add_heading('1.2.1 Kebutuhan Fungsional', level=3)
    
    kebutuhan_fungsional = [
        "Sistem mampu mencatat, mengubah, dan menghapus (CRUD) data warga terdampak yang diklasifikasikan ke dalam tipe Anak, Dewasa, dan Lansia dengan ID yang digenerasi secara otomatis.",
        "Sistem dapat menghitung skor prioritas evakuasi secara otomatis berdasarkan usia, jenis kelamin, kondisi kesehatan, dan tingkat bahaya lokasi secara dinamis (menggunakan konsep Polymorphism).",
        "Sistem mampu mengelola data posko (CRUD) dan memetakan warga ke posko pengungsian serta memantau sisa kapasitas dan stok logistik per wilayah secara real-time.",
        "Sistem mampu mencatat riwayat distribusi bantuan ke dalam file CSV serta memicu peringatan otomatis jika stok bantuan menipis atau kapasitas posko hampir penuh.",
        "Sistem mampu menghasilkan laporan rekapitulasi data (jumlah pengungsi, kebutuhan belum terpenuhi, sisa kapasitas per wilayah, total distribusi) yang dikunci (read-only) pada GUI dan dapat diekspor ke format TXT dan CSV secara terpisah."
    ]
    for req in kebutuhan_fungsional:
        add_numbered_list(doc, req)
        
    doc.add_heading('1.2.2 Kebutuhan Nonfungsional', level=3)
    kebutuhan_nonfungsional = [
        "Antarmuka aplikasi dibangun menggunakan Graphical User Interface (GUI) Tkinter agar mudah dioperasikan, dengan penyesuaian skala font (+5 point) untuk aksesibilitas pembacaan.",
        "Data disimpan secara persisten dalam format JSON (warga.json, posko.json) dan CSV (distribusi_bantuan.csv).",
        "Sistem wajib tangguh terhadap kesalahan input (resilient) dengan menangani kesalahan secara graceful melalui Custom Exception tanpa menyebabkan aplikasi terhenti (crash), serta melakukan penguncian input yang rentan typo melalui antarmuka dropdown (readonly)."
    ]
    for req in kebutuhan_nonfungsional:
        add_numbered_list(doc, req)

    doc.add_heading('1.3 Asumsi Sistem', level=2)
    doc.add_paragraph(
        "Sistem ini diasumsikan dijalankan oleh petugas atau koordinator relawan pada pusat komando bencana dengan akses ke perangkat komputer lokal. "
        "Data batas kritis kapasitas dan stok bantuan untuk setiap posko telah diperhitungkan secara manual oleh pakar di lapangan sebelum diinputkan ke dalam sistem."
    )
    
    doc.add_page_break()

    # --- BAB II ---
    doc.add_heading('BAB II: DESAIN PEMROGRAMAN BERORIENTASI OBJEK', level=1)
    
    doc.add_heading('2.1 UML Class Diagram', level=2)
    doc.add_paragraph("Struktur perangkat lunak dirancang menggunakan arsitektur yang memisahkan logika bisnis (Services), struktur data (Models), utilitas (Utils), dan antarmuka (GUI).")
    doc.add_paragraph("(*Catatan: Diagram UML disisipkan dari file docs/uml.md yang telah memuat relasi antara Subject, Observer, Manager, dan Kelas Warga*)")
    
    doc.add_heading('2.2 Hierarki Inheritance dan Bukti Polymorphism', level=2)
    doc.add_paragraph(
        "Konsep pewarisan (Inheritance) diterapkan secara jelas melalui kelas abstrak WargaTerdampak yang mewariskan atribut dasarnya "
        "(seperti nama, usia, jenis_kelamin) kepada tiga subclass: Anak, Dewasa, dan Lansia. Hal ini secara efektif menghilangkan redundansi deklarasi atribut umum pada tiap kelas."
    )
    doc.add_paragraph(
        "Untuk bukti Polymorphism, metode hitung_prioritas_evakuasi() di-override oleh masing-masing subclass. "
        "Saat sistem perlu mengurutkan antrean evakuasi, sistem cukup memanggil metode tersebut secara seragam dari kumpulan WargaTerdampak tanpa perlu menggunakan pernyataan pengecekan tipe beruntun. Objek Lansia akan mengembalikan skor dasar tertinggi (40), Anak menengah (30), "
        "dan Dewasa terendah (10)."
    )

    doc.add_heading('2.3 Mekanisme Encapsulation', level=2)
    doc.add_paragraph(
        "Integritas data dijaga secara ketat menggunakan mekanisme Encapsulation. Atribut __kondisi_kesehatan pada WargaTerdampak "
        "diubah menjadi atribut privat. Perubahan status hanya dapat dilakukan melalui property setter yang dilengkapi dengan logika validasi. "
        "Hal ini memastikan bahwa sistem hanya menerima nilai yang telah ditentukan. Jika input di luar batasan tersebut (bahkan dari modifikasi internal), sistem menolak perubahan."
    )

    doc.add_heading('2.4 Custom Exception', level=2)
    doc.add_paragraph(
        "Sistem ini mendefinisikan dua Custom Exception untuk memitigasi kesalahan logika bisnis yang spesifik:"
    )
    exceptions = [
        "DataKesehatanTidakValidError: Diturunkan dari ValueError. Dipicu saat enkapsulasi menolak input kondisi kesehatan yang tidak sesuai dengan standar sistem.",
        "KapasitasPoskoPenuhError: Diturunkan dari Exception. Dipicu saat terdapat upaya untuk memasukkan pengungsi ke posko yang nilai kapasitas terisinya sudah mencapai kapasitas maksimal."
    ]
    for exc in exceptions:
        p = doc.add_paragraph(exc, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_heading('2.5 Pola Desain (Observer Pattern)', level=2)
    doc.add_paragraph(
        "Sistem peringatan dini dibangun menggunakan Observer Pattern. Kelas Posko bertindak sebagai Subject, sementara "
        "kelas Relawan (dan komponen antarmuka GUI secara tidak langsung melalui inner class) bertindak sebagai Observer. Setiap kali terjadi penambahan/pengurangan pada kapasitas atau stok bantuan, "
        "Posko akan memvalidasi apakah ambang kritis telah tercapai. Jika ya, ia akan mengeksekusi metode notify() untuk mengirimkan broadcast "
        "peringatan kepada seluruh Observer yang berlangganan."
    )

    doc.add_heading('2.6 Prinsip SOLID (Single Responsibility Principle)', level=2)
    doc.add_paragraph(
        "Prinsip SRP diterapkan pada kelas DataManager. Kelas ini hanya memiliki satu alasan untuk berubah, yaitu terkait dengan mekanisme pembacaan dan penulisan penyimpanan persisten (File Handling JSON/CSV). "
        "DataManager sepenuhnya terisolasi dari logika perhitungan prioritas evakuasi maupun manipulasi antarmuka pengguna Tkinter."
    )
    
    doc.add_page_break()

    # --- BAB III ---
    doc.add_heading('BAB III: PENJELASAN IMPLEMENTASI FITUR UTAMA', level=1)
    doc.add_paragraph(
        "Bagian ini menjabarkan bagaimana lima konsep wajib diimplementasikan di dalam kode sumber Python."
    )
    
    doc.add_heading('3.1 Implementasi Inheritance & Polymorphism', level=2)
    doc.add_paragraph("Cuplikan kode pada src/models/warga.py yang menunjukkan metode abstrak yang ditimpa oleh subclass (Lansia dan Anak):")
    kode_poly = doc.add_paragraph()
    kode_poly.add_run(
        "class Lansia(WargaTerdampak):\n"
        "    def hitung_prioritas_evakuasi(self) -> int:\n"
        "        # Lansia mendapat base score tertinggi (40) + skor kondisi + skor bahaya\n"
        "        base_score = 40\n"
        "        return base_score + self._hitung_skor_kondisi() + self._hitung_skor_bahaya()\n\n"
        "class Anak(WargaTerdampak):\n"
        "    def hitung_prioritas_evakuasi(self) -> int:\n"
        "        base_score = 30\n"
        "        return base_score + self._hitung_skor_kondisi() + self._hitung_skor_bahaya()"
    )
    kode_poly.style = 'Intense Quote'
    
    doc.add_heading('3.2 Implementasi Encapsulation & Exception', level=2)
    doc.add_paragraph("Cuplikan property setter pada src/models/warga.py yang memvalidasi aliran data dan memicu custom exception saat diintervensi oleh nilai yang salah:")
    kode_encap = doc.add_paragraph()
    kode_encap.add_run(
        "    @kondisi_kesehatan.setter\n"
        "    def kondisi_kesehatan(self, value: str):\n"
        "        if value not in self.KONDISI_VALID:\n"
        "            raise DataKesehatanTidakValidError(value)\n"
        "        self.__kondisi_kesehatan = value"
    )
    kode_encap.style = 'Intense Quote'
    
    doc.add_heading('3.3 Implementasi File Handling & ID Auto-Increment', level=2)
    doc.add_paragraph(
        "Sistem menggunakan DataManager untuk membaca dan menulis daftar objek (yang telah diserialisasi ke dalam bentuk dictionary) ke dalam file JSON, serta log distribusi ke dalam CSV. "
        "Adapun untuk manajemen identitas (Primary Key), ID untuk posko dan warga dihasilkan secara otomatis (auto-increment) dengan mengekstraksi dan menambahkan nilai numerik dari string ID entitas terakhir yang tersimpan di dalam array memori."
    )

    doc.add_page_break()

    # --- BAB IV ---
    doc.add_heading('BAB IV: SKENARIO PENGUJIAN & EVALUASI', level=1)
    
    doc.add_heading('4.1 Tabel Hasil Pengujian Unit', level=2)
    doc.add_paragraph("Pengujian logika backend dilakukan menggunakan framework unittest bawaan Python dengan 8 skenario berbeda. Seluruh pengujian berjalan dengan sukses.")
    
    # Buat tabel sederhana di Word
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Modul'
    hdr_cells[1].text = 'Skenario'
    hdr_cells[2].text = 'Status'
    
    tests = [
        ("Polymorphism", "Eksekusi hitung_prioritas() pada 3 tipe objek; Lansia harus > Dewasa.", "Passed"),
        ("Encapsulation", "Set kondisi_kesehatan = 'Pusing' melempar DataKesehatanTidakValidError.", "Passed"),
        ("Posko Exception", "tambah_warga ke posko penuh melempar KapasitasPoskoPenuhError.", "Passed"),
        ("Observer", "Relawan menerima notifikasi saat batas kritis kapasitas posko tercapai.", "Passed"),
        ("File Handling", "Serialize ke warga.json lalu load kembali tanpa data yang hilang.", "Passed")
    ]
    
    for mod, sken, stat in tests:
        row_cells = table.add_row().cells
        row_cells[0].text = mod
        row_cells[1].text = sken
        row_cells[2].text = stat
        
    doc.add_paragraph('\n')
        
    doc.add_heading('4.2 Pengujian Kegagalan Operasi File (Edge Case)', level=2)
    doc.add_paragraph(
        "Untuk menguji keandalan sistem terhadap kegagalan operasi baca/tulis, skenario penghapusan file warga.json secara sengaja dilakukan saat program tidak berjalan. "
        "Saat program diinisialisasi ulang, alih-alih mengalami 'Crash' karena FileNotFoundError, blok try-except di dalam kelas DataManager menanganinya dengan senyap. "
        "Aplikasi kemudian mengembalikan list kosong (kembali ke state awal) sehingga pengguna tetap dapat melakukan pendataan dari awal dengan stabil."
    )
    
    doc.add_heading('4.3 Kendala dan Evaluasi Program', level=2)
    doc.add_paragraph(
        "Selama proses pengembangan, terdapat dua kendala utama yang dijumpai:"
    )
    kendala = [
        "Sinkronisasi Antarmuka: Ketika operasi Edit atau Hapus dilakukan, combobox pemilihan posko pada tab lain tidak serta-merta diperbarui. Masalah ini diselesaikan dengan menyematkan metode penyegaran (refresh) terpusat pada setiap eksekusi metode CRUD.",
        "Validasi Tipe Data pada Operasi Pencarian: Masalah tidak tertangkapnya ID warga saat proses penghapusan di GUI akibat ketidaksesuaian tipe data antara Integer dan String berhasil diatasi dengan penerapan konversi str() eksplisit pada perulangan filter."
    ]
    for k in kendala:
        add_numbered_list(doc, k)
    
    doc.add_heading('4.4 Kemungkinan Pengembangan', level=2)
    doc.add_paragraph(
        "Arsitektur file handling berbasis JSON murni pada DataManager sangat efisien untuk skalabilitas kecil hingga menengah. Namun, apabila sistem ini diimplementasikan untuk menangani puluhan ribu pengungsi lintas provinsi, "
        "kelas tersebut disarankan untuk ditransisikan menggunakan RDBMS seperti PostgreSQL. Selain itu, penambahan modul pemetaan geospasial dapat mempermudah relawan memantau kedekatan lokasi korban dengan posko terdekat."
    )
    
    doc.add_heading('4.5 Kesimpulan', level=2)
    doc.add_paragraph(
        "Aplikasi Sistem Manajemen Evakuasi Bencana dan Pengungsian ini telah dikembangkan dengan mematuhi seluruh spesifikasi tugas secara menyeluruh. "
        "Implementasi pola desain Observer sukses mewujudkan notifikasi rantai pasok secara otonom, sementara kombinasi Polymorphism, Encapsulation, dan Custom Exception terbukti "
        "secara fundamental memperkuat integritas data aplikasi, meminimalisir kesalahan operasional manusia, dan memudahkan proses pemeliharaan kode (maintainability) di masa mendatang."
    )

    doc.save('/home/darulqutni/Joki/jokiAksa/Makalah_Evakuasi_Bencana.docx')

if __name__ == '__main__':
    buat_makalah()
